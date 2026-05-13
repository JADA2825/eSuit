"""
reporte_docx.py — Reporte editable en Microsoft Word (.docx)

Genera el mismo contenido del PDF pero en formato Word para que el usuario
pueda editarlo libremente y agregar información adicional.

Tipografía:
  Cuerpo:    Garamond 12 pt, interlineado 1.5, justificado, negro
  Títulos:   Courier New 16 pt, azul
  Subtítulos: Courier New 14 pt, azul
  Fórmulas:  Courier New 11 pt, azul oscuro
"""
import io
import math
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from calculos import r_x_efectivas

# ═══════════════════════════════════════════════
# COLORES (mismos que el PDF)
# ═══════════════════════════════════════════════
COL_AZUL_TIT   = RGBColor(0x1F, 0x38, 0x64)
COL_AZUL_FORM  = RGBColor(0x0F, 0x2A, 0x5A)
COL_AZUL_BAR   = RGBColor(0x2D, 0x4A, 0x7A)
COL_NEGRO      = RGBColor(0x1A, 0x1A, 0x1A)
COL_SUBTXT     = RGBColor(0x5B, 0x5B, 0x5B)
COL_VERDE      = RGBColor(0x1D, 0x6E, 0x3A)
COL_ROJO       = RGBColor(0xA3, 0x16, 0x21)
COL_BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

# Sombreado de celdas
SHADE_HEADER   = "1F3864"   # azul oscuro encabezado
SHADE_ROW_ALT  = "F5F5F7"   # gris alternado
SHADE_RECOM    = "E7F5EA"   # verde claro recomendada


def _set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex="DCDCDC"):
    """Aplica borde sutil a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tc_pr.append(borders)


def _styled_run(par, text, *, font="Garamond", size=12, bold=False,
                italic=False, color=COL_NEGRO):
    """Agrega un run con tipografía explícita."""
    run = par.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Asegurar que Word use la fuente Garamond/Courier en todas las regiones (incl. asiático)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"),    font)
    rFonts.set(qn("w:hAnsi"),    font)
    rFonts.set(qn("w:cs"),       font)
    rFonts.set(qn("w:eastAsia"), font)
    return run


def _add_toc(doc, niveles="1-3"):
    """Inserta un campo Table of Contents nativo de Word.

    El TOC se rellena automáticamente cuando el usuario abre el .docx en Word
    (o presiona F9 para actualizarlo). Toma todos los párrafos con estilo
    Heading 1/2/3 que generamos con _add_titulo().
    """
    # Título visible "ÍNDICE"
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(8)
    _styled_run(p_title, "ÍNDICE", font="Courier New", size=16, bold=True,
                color=COL_AZUL_TIT)
    p_title.style = doc.styles["Heading 1"]
    for r in p_title.runs:
        r.font.name = "Courier New"; r.font.size = Pt(16)
        r.font.color.rgb = COL_AZUL_TIT; r.font.bold = True

    # Párrafo con el campo TOC
    p = doc.add_paragraph()
    run = p.add_run()
    rPr = run._element

    # fldChar w:fldCharType="begin"
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    rPr.append(fldChar_begin)

    # instrText: TOC \o "1-3" \h \z \u
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f' TOC \\o "{niveles}" \\h \\z \\u '
    rPr.append(instrText)

    # fldChar w:fldCharType="separate"
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    rPr.append(fldChar_sep)

    # Texto placeholder (visible si Word no actualiza el TOC automáticamente)
    placeholder = OxmlElement("w:t")
    placeholder.text = "Tabla de contenido (presiona F9 para actualizar en Word)"
    rPr.append(placeholder)

    # fldChar w:fldCharType="end"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    rPr.append(fldChar_end)

    # Indicar a Word que recalcule todos los campos al abrir el doc
    settings = doc.settings.element
    updateFields = settings.find(qn("w:updateFields"))
    if updateFields is None:
        updateFields = OxmlElement("w:updateFields")
        updateFields.set(qn("w:val"), "true")
        settings.append(updateFields)


def _add_titulo(doc, texto, nivel=1):
    """Título de sección (Courier New, azul). nivel: 1=16pt, 2=14pt, 3=12pt."""
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    _styled_run(p, texto, font="Courier New",
                size=sizes.get(nivel, 14),
                bold=True, color=COL_AZUL_TIT)
    # Marcar como heading para outline navigation (sin estilo)
    p.style = doc.styles[f"Heading {min(nivel, 9)}"]
    # Override del estilo de Heading
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(sizes.get(nivel, 14))
        run.font.color.rgb = COL_AZUL_TIT
        run.font.bold = True
    return p


def _add_cuerpo(doc, texto_html, justify=True):
    """Párrafo cuerpo: Garamond 12, interlineado 1.5, justificado, negro.
       Soporta <b>...</b> e <i>...</i> embebidos.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

    # Parser minimalista para <b> e <i>
    import re
    tokens = re.split(r"(<b>|</b>|<i>|</i>|<sub>|</sub>|<super>|</super>|<br/>|<br />)", texto_html)
    bold = italic = sub = sup = False
    for tk in tokens:
        if tk == "<b>": bold = True
        elif tk == "</b>": bold = False
        elif tk == "<i>": italic = True
        elif tk == "</i>": italic = False
        elif tk == "<sub>": sub = True
        elif tk == "</sub>": sub = False
        elif tk == "<super>": sup = True
        elif tk == "</super>": sup = False
        elif tk in ("<br/>", "<br />"):
            p.add_run().add_break()
        elif tk:
            r = _styled_run(p, tk, bold=bold, italic=italic)
            if sub:
                r.font.subscript = True
                r.font.size = Pt(9)
            if sup:
                r.font.superscript = True
                r.font.size = Pt(9)
    return p


def _add_formula(doc, texto):
    """Fórmula: Courier New 11 pt, azul oscuro, con sangría y separación."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    # Parsea <b> e <i>
    import re
    tokens = re.split(r"(<b>|</b>|<i>|</i>|<sub>|</sub>|<super>|</super>)", texto)
    bold = italic = sub = sup = False
    for tk in tokens:
        if tk == "<b>": bold = True
        elif tk == "</b>": bold = False
        elif tk == "<i>": italic = True
        elif tk == "</i>": italic = False
        elif tk == "<sub>": sub = True
        elif tk == "</sub>": sub = False
        elif tk == "<super>": sup = True
        elif tk == "</super>": sup = False
        elif tk:
            r = _styled_run(p, tk, font="Courier New", size=11,
                            bold=bold, italic=italic, color=COL_AZUL_FORM)
            if sub:
                r.font.subscript = True
                r.font.size = Pt(9)
            if sup:
                r.font.superscript = True
                r.font.size = Pt(9)
    return p


def _add_resultado(doc, texto, ok=True):
    """Mensaje de cumple / no cumple."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _styled_run(p, texto, font="Garamond", size=11, bold=True,
                color=COL_VERDE if ok else COL_ROJO)
    return p


def _add_tabla_datos(doc, filas, anchos_cm=(6.5, 10.5)):
    """Tabla 2 columnas: Parámetro | Valor."""
    table = doc.add_table(rows=1 + len(filas), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(anchos_cm[0])
    table.columns[1].width = Cm(anchos_cm[1])

    # Encabezado
    hdr = table.rows[0]
    for j, txt in enumerate(("Parámetro", "Valor")):
        cell = hdr.cells[j]
        cell.width = Cm(anchos_cm[j])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, SHADE_HEADER)
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(p, txt, font="Courier New", size=10, bold=True, color=COL_BLANCO)

    # Filas
    for i, (k, v) in enumerate(filas, start=1):
        row = table.rows[i]
        shade = SHADE_ROW_ALT if (i % 2 == 1) else None
        for j, (txt, font, bold, color) in enumerate(
            [(str(k), "Garamond", True, COL_AZUL_TIT),
             (str(v), "Garamond", False, COL_NEGRO)]
        ):
            cell = row.cells[j]
            cell.width = Cm(anchos_cm[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if shade:
                _set_cell_shading(cell, shade)
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            # Parsear <b>/<i>/<sub>/<super> en valor
            import re
            tokens = re.split(r"(<b>|</b>|<i>|</i>|<sub>|</sub>|<super>|</super>)", txt)
            b_open = bold
            i_open = False
            sub_open = sup_open = False
            for tk in tokens:
                if tk == "<b>": b_open = True
                elif tk == "</b>": b_open = bold
                elif tk == "<i>": i_open = True
                elif tk == "</i>": i_open = False
                elif tk == "<sub>": sub_open = True
                elif tk == "</sub>": sub_open = False
                elif tk == "<super>": sup_open = True
                elif tk == "</super>": sup_open = False
                elif tk:
                    r = _styled_run(p, tk, font=font, size=10,
                                    bold=b_open, italic=i_open, color=color)
                    if sub_open:
                        r.font.subscript = True; r.font.size = Pt(8)
                    if sup_open:
                        r.font.superscript = True; r.font.size = Pt(8)
    return table


def _add_tabla_candidatos(doc, df):
    """Tabla completa de conductores candidatos."""
    import pandas as _pd
    if isinstance(df, list):
        df = _pd.DataFrame(df) if df else None
    if df is None or len(df) == 0:
        return None
    # Sustitución de marcadores ya hecha por la app (Si/No/Rec.)
    # Pero por compatibilidad con DF que aún tenga ✓/✗/★:
    df = df.copy()
    df = df.replace({"✓": "Sí", "✗": "No", "★": "Rec."})

    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    table.autofit = True

    # Encabezado
    for j, c in enumerate(cols):
        cell = table.rows[0].cells[j]
        _set_cell_shading(cell, SHADE_HEADER)
        _set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(p, str(c), font="Courier New", size=9, bold=True, color=COL_BLANCO)

    # Filas
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        last_val = str(row.iloc[-1]) if len(row) else ""
        es_recom = "Rec." in last_val
        shade = SHADE_RECOM if es_recom else (SHADE_ROW_ALT if (i % 2 == 1) else None)
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            _set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if shade:
                _set_cell_shading(cell, shade)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _styled_run(p, str(val), font="Garamond", size=9,
                        bold=es_recom, color=COL_NEGRO)
    return table


def _add_tabla_tubos(doc, resultados_todos):
    """Tabla completa de tuberías."""
    cols = ["Tubería", "Ø int. (mm)", "Área int. (mm²)",
            "Área conds. (mm²)", "Relleno (%)", "Máx. %", "Cumple", "Sel."]
    table = doc.add_table(rows=1 + len(resultados_todos), cols=len(cols))
    table.autofit = True

    for j, c in enumerate(cols):
        cell = table.rows[0].cells[j]
        _set_cell_shading(cell, SHADE_HEADER)
        _set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(p, c, font="Courier New", size=9, bold=True, color=COL_BLANCO)

    for i, r in enumerate(resultados_todos, start=1):
        es_recom = r.get("recomendada", False)
        shade = SHADE_RECOM if es_recom else (SHADE_ROW_ALT if (i % 2 == 1) else None)
        diam = r.get("diametro", r.get("tubo", "—").split(" ")[0])
        vals = [
            str(r.get("tubo", "—")),
            str(diam),
            f"{r.get('area_tubo', 0):.1f}",
            f"{r.get('area_conds', 0):.2f}",
            f"{r.get('fill_pct', 0):.1f}",
            f"{r.get('fill_max', 40)}",
            "Sí" if r.get("cumple") else "No",
            "Rec." if es_recom else "",
        ]
        for j, v in enumerate(vals):
            cell = table.rows[i].cells[j]
            _set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if shade:
                _set_cell_shading(cell, shade)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _styled_run(p, v, font="Garamond", size=9,
                        bold=es_recom, color=COL_NEGRO)
    return table


def _add_tabla_resumen_circuitos(doc, circuitos):
    """Tabla resumen estilo tablero — todos los circuitos."""
    cols = ["No.", "Circuito", "P (W)", "V", "I (A)", "I.dis",
            "Conductor", "Amp", "CdT%", "Máx", "OCPD",
            "Tierra", "Tubería", "Estado"]
    table = doc.add_table(rows=1 + len(circuitos), cols=len(cols))
    table.autofit = True

    for j, c in enumerate(cols):
        cell = table.rows[0].cells[j]
        _set_cell_shading(cell, SHADE_HEADER)
        _set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(p, c, font="Courier New", size=8, bold=True, color=COL_BLANCO)

    for i, c in enumerate(circuitos, start=1):
        cdt = c.get("cdt", 0)
        cdt_max = c.get("cdt_max", 3.0)
        estado = "OK" if cdt <= cdt_max else "REVISAR"
        ci = c.get("conduit_info")
        tubo = ci["recomendada"]["tubo"] if (ci and ci.get("recomendada")) else "—"
        cond = f"{c.get('cf',1)}x{c.get('conductor','—')} {c.get('material','')}"
        tierra = f"{c.get('tierra_calibre','—')} {c.get('tierra_material','')}".strip()
        vals = [
            str(c.get("id", i)),
            str(c.get("nombre", f"Cto {i}")),
            f"{c.get('potencia',0):,}",
            str(c.get("voltaje", "—")),
            f"{c.get('corriente',0):.1f}",
            f"{c.get('corriente_diseño',0):.1f}",
            cond,
            f"{c.get('ampacity_corr',0):.0f}",
            f"{cdt:.2f}",
            f"{cdt_max}",
            str(c.get("proteccion_fmt", c.get("proteccion_A", "—"))),
            tierra,
            tubo,
            estado,
        ]
        shade = SHADE_ROW_ALT if (i % 2 == 1) else None
        for j, v in enumerate(vals):
            cell = table.rows[i].cells[j]
            _set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if shade:
                _set_cell_shading(cell, shade)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _styled_run(p, v, font="Garamond", size=8, color=COL_NEGRO)


# ═══════════════════════════════════════════════
# DESARROLLO DE CÁLCULO POR CIRCUITO
# ═══════════════════════════════════════════════
def _seccion_circuito(doc, circ, idx):
    nombre = circ.get("nombre", f"Circuito {idx}")
    mat = circ.get("material", "Cu")
    canal_lbl = circ.get("canalizacion_label", circ.get("canalizacion", "Acero"))
    config_code = circ.get("config_code", "mono2h")
    fp = circ.get("fp", 0.9)
    V = circ.get("voltaje", 120)
    P = circ.get("potencia", 0)
    I = circ.get("corriente", 0)
    fd = circ.get("factor_demanda", 1.25)
    I_dis = I * fd
    cf = circ.get("cf", 1)
    I_dis_cond = I_dis / cf
    cond = circ.get("conductor", "—")
    cdt = circ.get("cdt", 0)
    cdt_max = circ.get("cdt_max", 3.0)
    canal = circ.get("canalizacion", "AC")
    t_term = circ.get("temp_term", 75)

    _add_titulo(doc, f"7.{idx} {nombre}", nivel=2)

    # 7.1.1 Datos de entrada
    _add_titulo(doc, f"7.{idx}.1 Datos de entrada", nivel=3)
    filas_ent = [
        ("Potencia de la carga", f"P = {P:,} W"),
        ("Tensión de operación", f"V = {V} V"),
        ("Factor de potencia", f"cos φ = {fp}"),
        ("Configuración", str(circ.get("configuracion", "—"))),
        ("Longitud del circuito", f"L = {circ.get('longitud',0)} m"),
        ("Material conductor", f"{mat} ({'Cobre' if mat=='Cu' else 'Aluminio'})"),
        ("Conductores por fase", f"CF = {cf}"),
        ("Canalización (Tabla 9)", canal_lbl),
        ("Aislamiento del conductor", f"{circ.get('temp_aislamiento',75)} °C"),
        ("C.d.T. máxima permitida", f"{cdt_max} %"),
    ]
    _add_tabla_datos(doc, filas_ent)

    # 7.1.2 Cálculo de corriente
    _add_titulo(doc, f"7.{idx}.2 Cálculo de corriente nominal", nivel=3)
    if config_code == "trifasico":
        formula_I = "I = P / ( √3 · V · cosφ )"
        sustit = f"I = {P:,} / ( 1.7321 × {V} × {fp} ) = <b>{I:.3f} A</b>"
    else:
        formula_I = "I = P / ( V · cosφ )"
        sustit = f"I = {P:,} / ( {V} × {fp} ) = <b>{I:.3f} A</b>"
    _add_cuerpo(doc, "Fórmula aplicada:", justify=False)
    _add_formula(doc, formula_I)
    _add_cuerpo(doc, "Sustitución:", justify=False)
    _add_formula(doc, sustit)

    _add_cuerpo(doc, f"Corriente de diseño (factor de carga continua = {fd}):", justify=False)
    _add_formula(doc, f"I_diseño = I × FD = {I:.3f} × {fd} = <b>{I_dis:.3f} A</b>")
    if cf > 1:
        _add_formula(doc, f"I por conductor = I_diseño / CF = {I_dis:.3f} / {cf} = <b>{I_dis_cond:.3f} A</b>")

    # 7.1.3 Caída de tensión
    _add_titulo(doc, f"7.{idx}.3 Cálculo de caída de tensión", nivel=3)
    try:
        R_eff, X_eff = r_x_efectivas(cond, mat, canal, t_term)
        R75 = R_eff * (75 + 234.5) / (t_term + 234.5)
    except Exception:
        R_eff, X_eff, R75 = 0, 0, 0
    sen_phi = math.sqrt(max(0.0, 1 - fp**2))

    _add_cuerpo(doc,
        f"De la <b>Tabla 9 NOM-001</b>, para conductor <b>{cond} AWG {mat}</b> "
        f"en canalización <b>{canal_lbl}</b>:"
    )
    _add_formula(doc, f"R<sub>75</sub> = {R75:.4f} Ω/km · X = {X_eff:.4f} Ω/km")
    _add_cuerpo(doc, f"Ajuste a temperatura de terminal {t_term}°C (Art. 110-14):", justify=False)
    _add_formula(doc,
        f"R<sub>T</sub> = R<sub>75</sub> · (T + 234.5)/(75 + 234.5) = "
        f"{R75:.4f} × ({t_term}+234.5)/(75+234.5) = <b>{R_eff:.4f} Ω/km</b>"
    )

    if config_code == "trifasico":
        formula_cdt = "ΔV = √3 · I · (L/1000) · ( (R/CF)·cosφ + (X/CF)·senφ )"
        coef = "√3"
    else:
        formula_cdt = "ΔV = 2 · I · (L/1000) · ( (R/CF)·cosφ + (X/CF)·senφ )"
        coef = "2"

    L_km = circ.get("longitud", 0) / 1000
    Z = (R_eff / cf) * fp + (X_eff / cf) * sen_phi
    dV_v = (math.sqrt(3) if config_code == "trifasico" else 2) * I * L_km * Z
    pct = (dV_v / V * 100) if V else 0

    _add_cuerpo(doc, "Caída de tensión por el método de impedancia (Tabla 9):", justify=False)
    _add_formula(doc, formula_cdt)
    _add_cuerpo(doc, "Sustitución de valores:", justify=False)
    _add_formula(doc,
        f"ΔV = {coef} × {I:.3f} × ({circ.get('longitud',0)}/1000) × "
        f"( ({R_eff:.4f}/{cf}) × {fp} + ({X_eff:.4f}/{cf}) × {sen_phi:.4f} )"
    )
    _add_formula(doc, f"ΔV = <b>{dV_v:.4f} V</b>")
    if config_code == "trifasico":
        _add_formula(doc, "%ΔV = ΔV · 100 / V_línea")
    elif config_code == "mono2h":
        _add_formula(doc, "%ΔV = ΔV · 100 / V_fase-neutro")
    else:
        _add_formula(doc, "%ΔV = ΔV · 100 / V_línea")
    _add_formula(doc, f"%ΔV = ({dV_v:.4f} × 100) / {V} = <b>{pct:.3f} %</b>")

    cumple_cdt = cdt <= cdt_max
    _add_resultado(doc,
        f"Cumple el criterio de caída de tensión: %ΔV = {cdt:.2f}% ≤ {cdt_max}% (máx.)"
        if cumple_cdt else
        f"NO cumple: %ΔV = {cdt:.2f}% > {cdt_max}% (máx.). "
        f"Se requiere subir calibre o reducir longitud.",
        ok=cumple_cdt
    )

    # Tabla completa de candidatos
    tabla_cand = circ.get("tabla_conductor_df")
    if tabla_cand is not None:
        _add_cuerpo(doc,
            "<b>Tabla de verificación de conductores candidatos</b> "
            "— se evalúa cada calibre disponible del material. "
            "<i>Rec.</i> = recomendado (fila resaltada); "
            "<i>Sí/No</i> = cumple o no cumple cada criterio.",
            justify=False,
        )
        _add_tabla_candidatos(doc, tabla_cand)

    # 7.1.4 Protección
    _add_titulo(doc, f"7.{idx}.4 Protección contra sobrecorriente", nivel=3)
    ocpd_A = circ.get("proteccion_A", "—")
    ocpd_fmt = circ.get("proteccion_fmt", str(ocpd_A))
    polos = circ.get("polos", 1)
    _add_cuerpo(doc,
        "Según NOM-001 Art. 240 / NEC 240.6, se selecciona el siguiente tamaño "
        "estándar de protección que sea ≥ I_diseño/CF y no exceda la ampacidad "
        "corregida del conductor."
    )
    _add_formula(doc,
        f"I_diseño/CF = {I_dis_cond:.3f} A   →   Protección estándar siguiente: "
        f"<b>{ocpd_A} A</b>"
    )
    _add_formula(doc, f"En formato comercial ({polos} polo(s)): <b>{ocpd_fmt} A</b>")

    # 7.1.5 Tierra
    _add_titulo(doc, f"7.{idx}.5 Conductor de puesta a tierra (cable desnudo)", nivel=3)
    tierra_cal = circ.get("tierra_calibre", "—")
    tierra_mat = circ.get("tierra_material", "Cu")
    origen = "automático (Tabla 250-122 NOM-001)" if circ.get("tierra_auto", True) else "selección manual"
    _add_cuerpo(doc,
        f"Para una protección de <b>{ocpd_A} A</b>, la <b>Tabla 250-122</b> "
        f"de la NOM-001 especifica un conductor de puesta a tierra de equipos "
        f"calibre <b>{tierra_cal} AWG</b> de <b>{tierra_mat} desnudo</b> "
        f"({origen})."
    )

    # 7.1.6 Tubería
    ci = circ.get("conduit_info")
    if ci and ci.get("recomendada"):
        _add_titulo(doc, f"7.{idx}.6 Cálculo de tubería (conduit)", nivel=3)
        params = ci.get("params") or {}
        rec = ci.get("recomendada")
        _add_cuerpo(doc,
            f"Para <b>{params.get('num_conds','—')} conductor(es)</b> calibre "
            f"<b>{params.get('calibre','—')} AWG</b> con aislamiento "
            f"<b>{params.get('aislamiento','—')}</b>, en tubería tipo "
            f"<b>{params.get('tipo_tubo','—')}</b>:"
        )
        _add_tabla_tubos(doc, ci.get("resultados_todos") or [rec])
        cumple_t = rec.get("cumple", False)
        _add_resultado(doc,
            f"Tubería seleccionada: {rec.get('tubo','—')} · relleno "
            f"{rec.get('fill_pct',0):.1f}% ≤ {rec.get('fill_max',40)}%."
            if cumple_t else
            "Ningún tamaño cumple. Considerar tubería de mayor familia o más conductos.",
            ok=cumple_t
        )

    # 7.1.7 Resumen
    _add_titulo(doc, f"7.{idx}.7 Resumen", nivel=3)
    resumen = [
        ("Conductor seleccionado", f"{cf} × {cond} AWG {mat}"),
        ("Aislamiento", f"{circ.get('temp_aislamiento',75)} °C"),
        ("Ampacidad corregida", f"{circ.get('ampacity_corr',0):.1f} A"),
        ("Caída de tensión real", f"{cdt:.2f} % (máx. {cdt_max} %)"),
        ("Protección OCPD", f"{ocpd_fmt} A"),
        ("Conductor de tierra", f"{tierra_cal} AWG {tierra_mat} desnudo"),
    ]
    if ci and ci.get("recomendada"):
        resumen.append((
            "Tubería",
            f"{ci['recomendada'].get('tubo','—')} · "
            f"{ci['recomendada'].get('fill_pct',0):.1f}% relleno",
        ))
    _add_tabla_datos(doc, resumen)


# ═══════════════════════════════════════════════
# NORMATIVIDAD
# ═══════════════════════════════════════════════
NORMAS_DESCRIPCION = {
    "NOM-001-SEDE-2012": (
        "Norma Oficial Mexicana de Instalaciones Eléctricas. Documento de "
        "observancia obligatoria que establece las disposiciones para instalaciones "
        "eléctricas destinadas a la utilización de energía eléctrica en territorio nacional."
    ),
    "NOM-001-SEDE-2005": "Versión anterior de la NOM-001-SEDE.",
    "NEC 2023": "National Electrical Code (NFPA 70). Norma estadounidense referente.",
    "CFE DCDIAMT (Instalaciones aéreas MT)": (
        "Especificación CFE para la construcción de instalaciones eléctricas aéreas "
        "en media tensión, aplicable a redes de distribución y acometidas."
    ),
    "CFE DCDIASMT (Instalaciones subterráneas MT)": (
        "Especificación CFE para la construcción de instalaciones eléctricas "
        "subterráneas en media tensión: ductos, registros, cables y empalmes."
    ),
    "NOM-007-ENER-2014 (Eficiencia energética)": (
        "Eficiencia energética para sistemas de alumbrado en edificios no residenciales."
    ),
    "NMX-J-098-ANCE (Tensiones eléctricas estándar)": (
        "Valores normalizados de tensión eléctrica para sistemas de utilización."
    ),
    "NMX-J-235-ANCE (Conductores eléctricos)": (
        "Características y métodos de prueba para conductores eléctricos aislados."
    ),
    "IEEE Std 141 (Red Book)": (
        "Recommended Practice for Electric Power Distribution for Industrial Plants."
    ),
    "IEC 60364 (Instalaciones eléctricas en edificios)": (
        "Estándar internacional para instalaciones eléctricas en edificios."
    ),
}


# ═══════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ═══════════════════════════════════════════════
def generar_reporte_docx(datos_proyecto, resultado_conductor,
                          tabla_conductor=None,
                          resultado_conduit=None,
                          conduit_params=None,
                          circuitos=None) -> bytes:
    """Genera el reporte en formato .docx (editable en Word)."""
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    dp = datos_proyecto
    rc = resultado_conductor

    # Si single-circuit, fabricamos un solo elemento
    if circuitos is None:
        rc_circ = dict(rc) if rc else {}
        rc_circ.setdefault("id", 1)
        rc_circ.setdefault("nombre", "Circuito principal")
        if resultado_conduit and conduit_params:
            recom = next((x for x in resultado_conduit if x["cumple"]), resultado_conduit[-1])
            rc_circ["conduit_info"] = {
                "params": conduit_params,
                "recomendada": recom,
                "resultados_todos": resultado_conduit,
            }
        if tabla_conductor is not None:
            rc_circ["tabla_conductor_df"] = tabla_conductor
        circuitos_internos = [rc_circ]
    else:
        circuitos_internos = circuitos

    # ═══ PORTADA ═══
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    _styled_run(p, "MEMORIA TÉCNICA DE CÁLCULO\n",
                font="Courier New", size=22, bold=True, color=COL_AZUL_TIT)
    _styled_run(p, "Instalación Eléctrica\n",
                font="Garamond", size=14, italic=True, color=COL_AZUL_BAR)
    _styled_run(p, "Caída de tensión · Conductores · Protecciones · Puesta a tierra",
                font="Garamond", size=11, color=COL_SUBTXT)

    doc.add_paragraph()
    filas_portada = [
        ("Proyecto", dp.get("proyecto", "—")),
        ("Cliente / Propietario", dp.get("cliente", "—")),
        ("Ubicación", dp.get("ubicacion", "—")),
        ("Sistema eléctrico", dp.get("sistema", "—")),
        ("Norma principal", dp.get("norma", "NOM-001-SEDE-2012")),
        ("Temperatura ambiente", f"{dp.get('temp_ambiente', 35)} °C"),
        ("Responsable técnico", dp.get("responsable", "—")),
        ("Cédula profesional", dp.get("cedula", "—")),
        ("Fecha de elaboración", str(dp.get("fecha", date.today()))),
    ]
    _add_tabla_datos(doc, filas_portada)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    _styled_run(p, "Documento técnico generado con eSuit · v2.0",
                font="Garamond", size=9, italic=True, color=COL_SUBTXT)

    doc.add_page_break()

    # ═══ ÍNDICE (TOC nativo de Word — se actualiza con F9 o al abrir) ═══
    _add_toc(doc, niveles="1-3")
    doc.add_page_break()

    # ═══ 1. INTRODUCCIÓN ═══
    _add_titulo(doc, "1. INTRODUCCIÓN", nivel=1)
    _add_cuerpo(doc,
        f"La presente memoria técnica documenta el diseño y cálculo de la instalación "
        f"eléctrica del proyecto <b>{dp.get('proyecto','—')}</b>, conforme a la normativa "
        f"mexicana vigente <b>{dp.get('norma','NOM-001-SEDE-2012')}</b>. El documento "
        f"describe los criterios de diseño, parámetros de entrada, fórmulas utilizadas "
        f"y desarrollo detallado de los cálculos de caída de tensión, selección de "
        f"conductores, dimensionamiento de protecciones contra sobrecorriente y "
        f"conductores de puesta a tierra. El propósito es servir como soporte técnico "
        f"para la construcción, supervisión y verificación de la instalación, "
        f"garantizando el cumplimiento de los requisitos mínimos de seguridad "
        f"eléctrica establecidos por la norma."
    )

    # ═══ 2. OBJETIVO ═══
    _add_titulo(doc, "2. OBJETIVO", nivel=1)
    _add_cuerpo(doc,
        "Determinar, mediante cálculo justificado, los siguientes elementos para "
        "cada uno de los circuitos del proyecto:"
    )
    items = [
        "<b>a)</b> Corriente de carga y corriente de diseño (incluyendo factor 1.25 si aplica).",
        "<b>b)</b> Calibre del conductor por criterio de ampacidad corregida.",
        "<b>c)</b> Verificación de la caída de tensión (% C.d.T.) dentro del límite normativo.",
        "<b>d)</b> Capacidad nominal de la protección contra sobrecorriente (OCPD) en formato comercial NxA.",
        "<b>e)</b> Calibre del conductor de puesta a tierra de equipos según NOM-001 Tabla 250-122.",
        "<b>f)</b> Tubería (conduit) suficiente para alojar los conductores cumpliendo el % de relleno permitido.",
    ]
    for it in items:
        _add_cuerpo(doc, it)

    # ═══ 3. DESCRIPCIÓN ═══
    _add_titulo(doc, "3. DESCRIPCIÓN DEL PROYECTO", nivel=1)
    desc = dp.get("descripcion") or (
        f"Instalación eléctrica para {dp.get('proyecto','—')} localizada en "
        f"{dp.get('ubicacion','—')}, propiedad de {dp.get('cliente','—')}. "
        f"El sistema de alimentación principal es {dp.get('sistema','—')}, "
        f"con temperatura ambiente de diseño de {dp.get('temp_ambiente',35)}°C."
    )
    _add_cuerpo(doc, desc)

    _add_cuerpo(doc, "<b>Datos generales</b>", justify=False)
    filas = [
        ("Proyecto", dp.get("proyecto", "—")),
        ("Cliente / Propietario", dp.get("cliente", "—")),
        ("Ubicación", dp.get("ubicacion", "—")),
        ("Sistema eléctrico principal", dp.get("sistema", "—")),
        ("Temperatura ambiente", f"{dp.get('temp_ambiente', 35)} °C"),
        ("Responsable técnico", dp.get("responsable", "—")),
        ("Cédula profesional", dp.get("cedula", "—")),
        ("Fecha", str(dp.get("fecha", date.today()))),
    ]
    _add_tabla_datos(doc, filas)

    doc.add_page_break()

    # ═══ 4. NORMATIVIDAD ═══
    _add_titulo(doc, "4. NORMATIVIDAD APLICABLE", nivel=1)
    _add_cuerpo(doc,
        "El presente cálculo se apega a las siguientes normas técnicas y especificaciones:"
    )
    principales = [dp.get("norma", "NOM-001-SEDE-2012")] + list(dp.get("normas_extra", []))
    seen = set()
    seleccionadas = [n for n in principales if not (n in seen or seen.add(n))]
    for i, n in enumerate(seleccionadas, 1):
        _add_cuerpo(doc,
            f"<b>{i}. {n}.</b> "
            + NORMAS_DESCRIPCION.get(n, "Norma técnica de referencia.")
        )

    # ═══ 5. CRITERIOS DE DISEÑO ═══
    _add_titulo(doc, "5. CRITERIOS DE DISEÑO", nivel=1)
    _add_cuerpo(doc,
        "La selección del conductor de cada circuito se realiza verificando "
        "simultáneamente dos criterios:"
    )
    _add_cuerpo(doc,
        "<b>(1) Capacidad de conducción (ampacidad).</b> La ampacidad base se "
        "toma de la <b>Tabla 310-15(B)(16)</b> de la NOM-001 según el material "
        "del conductor y la temperatura del aislamiento. Se afecta por el factor "
        "de corrección por temperatura ambiente y por el factor de agrupamiento. "
        "Debe satisfacer <i>Ampacidad ≥ I_diseño / CF</i>."
    )
    _add_cuerpo(doc,
        "<b>(2) Caída de tensión.</b> Se calcula con los valores de resistencia "
        "(R) y reactancia (X) de la <b>Tabla 9</b> de la NOM-001, ajustados a la "
        "temperatura del terminal del equipo según el Artículo 110-14. Para "
        "corrientes ≤ 100 A se asume terminal de 60 °C; para corrientes > 100 A "
        "se asume 75 °C; para motores (excepto Clase A) se usa 75 °C."
    )
    _add_cuerpo(doc,
        "<b>(3) Protección contra sobrecorriente.</b> Se selecciona el siguiente "
        "tamaño estándar (NEC 240.6 / NOM-001 Art. 240) que sea ≥ I_diseño y que "
        "no exceda la ampacidad corregida del conductor. El formato comercial se "
        "expresa como N x A donde N es el número de polos del sistema "
        "(1 monofásico, 2 bifásico, 3 trifásico)."
    )
    _add_cuerpo(doc,
        "<b>(4) Conductor de puesta a tierra de equipos.</b> Se dimensiona "
        "según la <b>Tabla 250-122</b> de la NOM-001, en función de la capacidad "
        "nominal de la protección contra sobrecorriente del circuito."
    )

    _add_cuerpo(doc, "<b>Constantes y parámetros base</b>", justify=False)
    filas = [
        ("Base de ampacidad (T<sub>0</sub>)", "30 °C"),
        ("Temperatura ambiente de proyecto", f"{dp.get('temp_ambiente',35)} °C"),
        ("Coeficiente α térmico Cu (R)", "α = 1/234.5 K<super>-1</super>"),
        ("Ajuste de R a temp. terminal",
         "R<sub>T</sub> = R<sub>75</sub> × (T + 234.5) / (75 + 234.5)"),
        ("Factor de carga continua", "1.25 (cargas ≥ 3 h continuas, NEC 210.20)"),
    ]
    _add_tabla_datos(doc, filas)

    doc.add_page_break()

    # ═══ 6. TABLA RESUMEN ═══
    _add_titulo(doc, "6. TABLA RESUMEN DE CIRCUITOS", nivel=1)
    _add_cuerpo(doc,
        f"Resumen de los <b>{len(circuitos_internos)} circuito(s)</b> del proyecto. "
        "Cada columna se desarrolla en detalle en la Sección 7."
    )
    _add_tabla_resumen_circuitos(doc, circuitos_internos)

    doc.add_page_break()

    # ═══ 7. MEMORIA POR CIRCUITO ═══
    _add_titulo(doc, "7. MEMORIA DE CÁLCULO POR CIRCUITO", nivel=1)
    _add_cuerpo(doc,
        "A continuación se presenta el desarrollo paso a paso de cada circuito: "
        "datos de entrada, cálculo de corriente, caída de tensión por método de "
        "impedancia (Tabla 9 NOM con ajuste por temperatura de terminal), "
        "selección de protección, conductor de tierra y tubería."
    )
    for i, c in enumerate(circuitos_internos, 1):
        _seccion_circuito(doc, c, i)
        if i < len(circuitos_internos):
            p = doc.add_paragraph()
            _styled_run(p, "_" * 80, font="Garamond", size=8, color=COL_SUBTXT)

    doc.add_page_break()

    # ═══ 8. CONCLUSIONES ═══
    _add_titulo(doc, "8. CONCLUSIONES", nivel=1)
    total_kw = sum(c.get("potencia", 0) for c in circuitos_internos) / 1000
    ok_count = sum(1 for c in circuitos_internos
                   if c.get("cdt", 0) <= c.get("cdt_max", 3.0))
    _add_cuerpo(doc,
        f"<b>1.</b> El proyecto comprende <b>{len(circuitos_internos)} circuito(s)</b> "
        f"con una carga total instalada de <b>{total_kw:.2f} kW</b>."
    )
    _add_cuerpo(doc,
        f"<b>2.</b> <b>{ok_count} de {len(circuitos_internos)}</b> circuito(s) "
        f"cumplen simultáneamente con los criterios de ampacidad y caída de "
        f"tensión establecidos en la {dp.get('norma','NOM-001-SEDE-2012')}."
    )
    _add_cuerpo(doc,
        "<b>3.</b> Los cálculos de resistencia incluyen el ajuste por temperatura "
        "del terminal del equipo conforme al Artículo 110-14 de la NOM-001-SEDE, "
        "y los valores de R y X corresponden a la Tabla 9 de la misma norma."
    )
    _add_cuerpo(doc,
        "<b>4.</b> Las protecciones contra sobrecorriente fueron seleccionadas "
        "como los tamaños estándar comerciales (NEC 240.6) inmediatamente "
        "superiores a la corriente de diseño, dentro de la ampacidad corregida "
        "del conductor."
    )
    _add_cuerpo(doc,
        "<b>5.</b> Los conductores de puesta a tierra de equipos fueron "
        "dimensionados según la Tabla 250-122 de la NOM-001 en función de la "
        "capacidad nominal del OCPD, utilizando cable desnudo del material indicado."
    )
    if ok_count < len(circuitos_internos):
        _add_resultado(doc,
            "Observación: Los circuitos marcados como REVISAR requieren ajuste "
            "(subir calibre, reducir longitud o redistribuir carga) antes de su "
            "construcción.",
            ok=False
        )

    # ═══ 9. FIRMA ═══
    doc.add_paragraph()
    _add_titulo(doc, "9. FIRMA DEL RESPONSABLE TÉCNICO", nivel=1)
    doc.add_paragraph()
    p = doc.add_paragraph("_" * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nombre = dp.get("responsable", "_____________________________")
    cedula = dp.get("cedula", "")
    fecha_doc = str(dp.get("fecha", date.today()))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, nombre, font="Garamond", size=11, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p2, "Responsable Técnico", font="Garamond", size=9, color=COL_SUBTXT)
    if cedula:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(p3, f"Cédula Profesional: {cedula}",
                    font="Garamond", size=9, color=COL_SUBTXT)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p4, fecha_doc, font="Garamond", size=10, color=COL_SUBTXT)

    # ═══ Guardar a bytes ═══
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
